from __future__ import annotations

from pathlib import Path

from docxtpl import DocxTemplate
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt



def generate_docx(template_path: Path, context: dict, output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = DocxTemplate(str(template_path))
    doc.render(context)
    doc.save(str(output_path))


def _set_font(paragraph, size: int = 9, bold: bool = False) -> None:
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold


def _cell_text(cell, text: str, *, size: int = 9, bold: bool = False, align=None) -> None:
    cell.text = str(text or "")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        if align is not None:
            paragraph.alignment = align
        _set_font(paragraph, size=size, bold=bold)


def _add_labeled_line(document: Document, label: str, value: str, note: str = "") -> None:
    table = document.add_table(rows=2, cols=2)
    table.autofit = True
    table.cell(0, 0).width = Cm(4)
    table.cell(0, 1).width = Cm(22)
    _cell_text(table.cell(0, 0), label, size=8)
    _cell_text(table.cell(0, 1), value, size=10)
    _cell_text(table.cell(1, 0), "", size=7)
    _cell_text(table.cell(1, 1), note, size=7, align=WD_ALIGN_PARAGRAPH.CENTER)


def generate_limit_card_docx(context: dict, output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)

    header = doc.add_table(rows=1, cols=3)
    header.autofit = True
    _cell_text(header.cell(0, 0), "Сільгосппідприємство\n" + context["company"], size=8)
    _cell_text(
        header.cell(0, 1),
        "Лімітно-забірна картка N ___\nна отримання матеріальних цінностей\nза "
        + context["date"],
        size=9,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _cell_text(
        header.cell(0, 2),
        "ЗАТВЕРДЖУЮ\nКерівник ____________\n\"___\" __________ 20__ р.",
        size=8,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    _add_labeled_line(doc, "Структурний підрозділ", context["structural_unit"])
    _add_labeled_line(
        doc,
        "Одержувач",
        context["recipient_line"],
        "трактор / автомобіль, державний номер, ПІБ одержувача",
    )
    _add_labeled_line(doc, "", context["recipient_name"], "ПІБ одержувача")

    table = doc.add_table(rows=1, cols=10)
    table.style = "Table Grid"
    headings = [
        "Дата",
        "Найменування цінностей",
        "Од. виміру",
        "Ліміт",
        "Кількість видано",
        "Залишок",
        "Підпис одержувача",
        "Разом видано",
        "Всього з урах. повернутого",
        "Бухгалтер",
    ]
    for index, heading in enumerate(headings):
        _cell_text(table.cell(0, index), heading, size=7, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for row_data in context["rows"]:
        row = table.add_row()
        values = [
            row_data["date"],
            row_data["fuel"],
            row_data["unit"],
            row_data["limit"],
            row_data["issued"],
            row_data["balance"],
            context["driver_signature"],
            "",
            "",
            "",
        ]
        for index, value in enumerate(values):
            _cell_text(row.cells[index], value, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    summary = table.add_row().cells
    _cell_text(summary[0], "Разом", size=8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(summary[4], context["total_issued"], size=8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(summary[7], context["total_issued"], size=8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(summary[8], context["total_with_returned"], size=8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(summary[9], context["accountant"], size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    for _ in range(max(0, 8 - len(context["rows"]))):
        row = table.add_row()
        for cell in row.cells:
            _cell_text(cell, "", size=8)

    signatures = doc.add_paragraph()
    signatures.add_run("\nВідпустив ____________    Одержав ____________ ")
    signatures.add_run(context["driver_signature"])
    signatures.add_run("    Бухгалтер ____________ ")
    signatures.add_run(context["accountant"])
    _set_font(signatures, size=9)

    doc.save(str(output_path))
